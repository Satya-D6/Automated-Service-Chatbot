from flask import Flask, request, jsonify, send_file, render_template
from fpdf import FPDF
import os
from datetime import datetime
import logging
from docx import Document

app = Flask(__name__)

# -------------------------------
# Folder Configuration
# -------------------------------

TEMPLATE_FOLDER = "templates"
GENERATED_FOLDER = "generated"

os.makedirs(GENERATED_FOLDER, exist_ok=True)

# -------------------------------
# Template Mapping
# -------------------------------

template_mapping = {
    "bonafide": "bonafide_template.docx",
    "study_certificate": "study_template.docx",
    "fee_receipt": "fee_receipt_template.docx",
    "tc": "transfer_template.docx",
    "conduct" : "conduct_template.docx"
}

# -------------------------------
# Home Route
# -------------------------------

@app.route('/')
def home():
    return "AI Certificate Generation API Running"

# -------------------------------
# Generate Certificate API
# -------------------------------

# Enhanced function to replace placeholders in the DOCX template
def replace_placeholders(template_path, data):
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"  # Ensure placeholders are in the format {{key}}
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    for table in doc.tables:  # Handle placeholders in tables
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))

    return doc

@app.route('/generate-certificate', methods=['POST'])
def generate_certificate():

    try:
        data = request.json

        # Get Request Data
        certificate_type = data.get("certificate_type")
        if certificate_type not in template_mapping:
            return jsonify({"status": "error", "message": "Invalid certificate type"}), 400

        template_path = os.path.join(TEMPLATE_FOLDER, template_mapping[certificate_type])
        if not os.path.exists(template_path):
            return jsonify({"status": "error", "message": f"Template file not found at '{template_path}'"}), 404

        # Replace placeholders in the template
        filled_doc = replace_placeholders(template_path, data)

        # Save the filled DOCX as the final output
        output_docx_path = os.path.join(GENERATED_FOLDER, f"{data['student_name']}_{certificate_type}.docx")
        filled_doc.save(output_docx_path)

        # Return the DOCX file as the response
        return send_file(output_docx_path, as_attachment=True)

    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

# -------------------------------
# Serve HTML Frontend
# -------------------------------

@app.route('/form')
def form():
    return render_template('index.html')

# -------------------------------
# Extract Placeholders API
# -------------------------------

def extract_placeholders(template_path):
    """Extract placeholders from a DOCX template."""
    doc = Document(template_path)
    placeholders = set()

    # Extract placeholders from paragraphs
    for paragraph in doc.paragraphs:
        for word in paragraph.text.split():
            if word.startswith("{{") and word.endswith("}}"):
                placeholders.add(word.strip("{}"))

    # Extract placeholders from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for word in cell.text.split():
                    if word.startswith("{{") and word.endswith("}}"):
                        placeholders.add(word.strip("{}"))

    return list(placeholders)

# Example usage to extract placeholders from all templates
@app.route('/extract-placeholders', methods=['GET'])
def extract_all_placeholders():
    try:
        all_placeholders = {}
        for cert_type, template_file in template_mapping.items():
            template_path = os.path.join(TEMPLATE_FOLDER, template_file)
            all_placeholders[cert_type] = extract_placeholders(template_path)

        return jsonify(all_placeholders)
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

# -------------------------------
# Run Flask App
# -------------------------------

if __name__ == '__main__':
    app.run(debug=True)