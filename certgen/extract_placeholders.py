from docx import Document

def extract_placeholders(docx_path):
    doc = Document(docx_path)
    placeholders = set()

    for paragraph in doc.paragraphs:
        if "{{" in paragraph.text and "}}" in paragraph.text:
            placeholders.update(
                part.strip("{}") for part in paragraph.text.split() if "{{" in part and "}}" in part
            )

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{{" in cell.text and "}}" in cell.text:
                    placeholders.update(
                        part.strip("{}") for part in cell.text.split() if "{{" in part and "}}" in part
                    )

    return placeholders

if __name__ == "__main__":
    template_path = "templates/conduct_template.docx"
    placeholders = extract_placeholders(template_path)
    print("Extracted placeholders:", placeholders)